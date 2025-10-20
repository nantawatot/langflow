# from langflow.field_typing import Data
import base64
import json
import os
import re

from docx import Document
from langflow.custom import Component
from langflow.io import MessageTextInput, Output
from langflow.schema import Message


class GetDocumentTemplate(Component):
    display_name = "Get Document Template"
    description = "Use this to retrieve the template document (base64), based on the Flow Profile"
    documentation: str = "https://docs.langflow.org/components-custom-components"
    icon = "Omniscien"
    name = "GetDocumentTemplate"

    inputs = [
        MessageTextInput(
            name="doc_id_key",
            display_name="Doc ID Key",
            info="This is a component to retrieve a document template based on the flow ID.",
            tool_mode=True,
            advanced=True,
        ),
        MessageTextInput(
            name="flow_id_key",
            display_name="Flow ID Key",
            info="This is a component to retrieve a document template based on the flow ID.",
            tool_mode=True,
            advanced=True,
        ),
        MessageTextInput(
            name="job_payload",
            display_name="Job Payload",
            info="This is a component to retrieve a document template based on the flow ID.",
            tool_mode=True,
        ),
        MessageTextInput(
            name="flow_profile",
            display_name="Flow Profile",
            info="This is a component to retrieve a document template based on the flow ID.",
            tool_mode=True,
        ),
        MessageTextInput(
            name="doc_template_id_to_merge",
            display_name="Doc Template Id to Merge",
            info="Choose which Doc Template Id to Merge.",
            tool_mode=True,
            advanced=True,
        ),
    ]

    outputs = [
        Output(display_name="Output", name="output", method="build_output"),
    ]

    def build_output(self) -> Message:
        # To obtain the flow ID and document ID from the input
        # flow_profile = self.obtain_key_value(self.flow_profile, self.doc_id_key)
        # document_id = self.obtain_document_ids(self.flow_profile)
        flow_id = self.obtain_key_value(self.job_payload, self.flow_id_key)

        # Condition: If doc_template_id_to_merge is provided, use it; otherwise, use the first document ID from the list
        if self.doc_template_id_to_merge and self.doc_template_id_to_merge.strip():
            document_id = self.doc_template_id_to_merge
        else:
            document_id = self.obtain_document_ids(self.flow_profile)[0]

        self.log(f"Document ID: {document_id}")

        api_url = "http://devdemo.languagestudio.com:4000/lsrestapi/v6/genai/getflowdoctemplate"

        params = {
            "id": flow_id,  # self.flow_id will be the value from the input
            "doctemplateid": document_id,
        }

        response = requests.get(api_url, params=params)

        self.log(f"Response from API: {response.text}")

        if response.status_code != 200:
            self.log(f"Error: {response.status_code} - {response.text}")
            raise Exception(f"Failed to retrieve document template: {response.text}")

        document_path = self.decode_base64_to_docx(response.text)
        self.log(f"Document path: {document_path}")

        styles = self.get_placeholder_styles(document_path.text)
        self.log(f"Styles Extracted: {styles}")

        styles_json = json.dumps(styles, indent=4)
        self.log(f"Styles JSON: {styles_json}")

        # Store styles in the graph context for later use
        try:
            parsed_json = json.loads(styles_json)
            self.graph.context["placeholder_styles"] = parsed_json
        except json.JSONDecodeError:
            pass

        return Message(text=str(document_path.text))

    def obtain_key_value(self, payload: str, key: str) -> str:
        # Extract the string from Message or use as-is
        input_str = payload.text if hasattr(payload, "text") else payload

        # First load
        data = json.loads(input_str)

        # If it's still a string, load again (double-encoded case)
        if isinstance(data, str):
            data = json.loads(data)

        # Navigate the nested keys
        keys = key.split(".")
        for k in keys:
            data = data[k]

        # Convert to string and remove unnecessary characters
        result_str = str(data).replace("[", "").replace("]", "").replace("'", "").replace('"', "")

        return result_str

    def obtain_document_ids(self, payload: str) -> list:
        def find_ids(obj):
            ids = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    if k == "id":
                        ids.append(v)
                    ids.extend(find_ids(v))
            elif isinstance(obj, list):
                for item in obj:
                    ids.extend(find_ids(item))
            return ids

        data = json.loads(payload)
        doctemplates = data["result"]["doctemplates"]["doctemplate"]
        return [item["id"] for item in doctemplates]

    def decode_base64_to_docx(self, input_value: str) -> Message:
        json_string_from_input = input_value
        self.log(f"Input to decode_base64_to_docx: {json_string_from_input}")

        try:
            # Step 1: Parse the outer JSON
            outer_data = json.loads(json_string_from_input)

            self.log(f"Successfully parsed outer JSON: {outer_data}")

            # The 'value' field is itself a JSON string that needs to be parsed
            inner_json_string = outer_data["result"][0]["doctemplate"]
            # inner_data = json.loads(inner_json_string)
            self.log(f"Successfully parsed inner JSON: {inner_json_string}")

            # Step 2: Extract the doctemplate payload
            # The 'result' is a list, and the doctemplate is inside the first item
            doctemplate_base64 = inner_json_string
            print("Successfully extracted Base64 string.")
            self.log("Successfully extracted Base64 string.")

            # Step 3: Base64 Decode the doctemplate string
            # It's important to convert the string to bytes before decoding
            decoded_bytes = base64.b64decode(doctemplate_base64)
            print("Successfully Base64 decoded the string.")
            self.log("Successfully Base64 decoded the string.")

            # Step 4: Save the decoded bytes as a .docx file
            output_filename = "output_document.docx"
            with open(output_filename, "wb") as f:
                f.write(decoded_bytes)

            print(f"Word document saved successfully as '{output_filename}'")
            self.log(f"Word document saved successfully as '{output_filename}'")

            current_directory = os.getcwd()

            # Log Path being saved to
            path = os.path.join(current_directory, output_filename)  # Return the full path

        except json.JSONDecodeError as e:
            print(f"Error decoding JSON: {e}")
            self.log(f"Error decoding JSON: {e}")
        except KeyError as e:
            print(f"Key not found in JSON payload: {e}")
            self.log(f"Key not found in JSON payload: {e}")
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            self.log(f"An unexpected error occurred: {e}")
        self.log(f"Returning path: {path}")
        return Message(text=str(path))

    def find_all_placeholders(self, docx_path: str):
        """Opens a docx file and finds all placeholders matching the [[...]] pattern.
        Returns a unique list of all found placeholders.
        """
        found_placeholders = []
        placeholder_pattern = re.compile(r"\[\[.*?\]\]")

        try:
            document = Document(docx_path)

            # Check all paragraphs in the main body
            for para in document.paragraphs:
                matches = placeholder_pattern.findall(para.text)
                found_placeholders.extend(matches)

            # Check all cells within all tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = placeholder_pattern.findall(cell.text)
                        found_placeholders.extend(matches)

            # Return a list of unique placeholders
            return sorted(list(set(found_placeholders)))

        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    def get_placeholder_styles(self, docx_path: str):
        """Opens the docx file, finds all placeholders, and extracts their font styles,
        handling cases where placeholders are split across multiple runs.
        """
        document = Document(docx_path)
        placeholder_info = {}
        placeholder_pattern = re.compile(r"\[\[.*?\]\]")

        def process_element_for_styles(element):
            combined_text = ""
            run_map = []

            # Step A: Build a map of character indices to their runs
            for run in element.runs:
                start_index = len(combined_text)
                combined_text += run.text
                run_map.append({"run": run, "start_index": start_index})

            # Step B: Find all placeholders in the combined text
            matches = placeholder_pattern.finditer(combined_text)

            # Step C: For each match, find the corresponding run and get the style
            for match in matches:
                placeholder_text = match.group(0)

                # Check if placeholder is already processed (to avoid duplicates)
                # Remove brackets from the placeholder name for the key
                clean_placeholder_text = placeholder_text.strip("[]")
                if clean_placeholder_text not in placeholder_info:
                    start_index = match.start()

                    # Find the run that contains the beginning of the placeholder
                    target_run = None
                    for item in run_map:
                        if item["start_index"] <= start_index < item["start_index"] + len(item["run"].text):
                            target_run = item["run"]
                            break

                    if target_run:
                        # Get font attributes with necessary fallbacks
                        run_font = target_run.font
                        run_style_font = target_run.style.font if target_run.style else None
                        para_style_font = element.style.font if element.style else None

                        font_name = (
                            run_font.name
                            or (run_style_font.name if run_style_font else None)
                            or (para_style_font.name if para_style_font else None)
                        )
                        font_size = (
                            run_font.size
                            or (run_style_font.size if run_style_font else None)
                            or (para_style_font.size if para_style_font else None)
                        )

                        placeholder_info[clean_placeholder_text] = {
                            "placeholder": clean_placeholder_text,
                            "font_name": font_name,
                            "font_size": font_size.pt if font_size else None,
                            "bold": target_run.bold,
                            "italic": target_run.italic,
                            "underline": target_run.underline,
                            "paragraph_style": element.style.name if element.style else None,
                        }

        # Process all paragraphs in the main document body
        for paragraph in document.paragraphs:
            process_element_for_styles(paragraph)

        # Process all paragraphs within tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_element_for_styles(paragraph)

        return list(placeholder_info.values())
